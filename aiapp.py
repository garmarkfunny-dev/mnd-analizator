# app.py
# -*- coding: utf-8 -*-
import os
import io
import json
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple

import streamlit as st
import pandas as pd
import altair as alt

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

APP_TITLE = "МНД — Анализатор"
HISTORY_JSON = "history.json"
SETTINGS_JSON = "settings.json"

DEFAULT_SETTINGS: Dict[str, Any] = {
    "theme": "dark",
    "accent_color": "#2b6cdf",
    "title_color":  "#2b6cdf",
    "embed_chart": True,
    "autosave_history": True,
}

def load_json(path: str, fallback):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return fallback

def save_json(path: str, data):
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def safe_float(x):
    try:
        if x is None or (isinstance(x, str) and not x.strip()):
            return None
        return float(str(x).replace(",", "."))
    except Exception:
        return None

def clamp_int(x: float) -> int:
    try:
        return int(max(0, min(100, round(x))))
    except Exception:
        return 0

def ensure_state_defaults():
    if "settings" not in st.session_state:
        st.session_state.settings = load_json(SETTINGS_JSON, DEFAULT_SETTINGS.copy())
    st.session_state.setdefault("pid", "")
    st.session_state.setdefault("tsh", "")
    st.session_state.setdefault("ft4", "")
    st.session_state.setdefault("ft3", "")
    st.session_state.setdefault("at_tpo", "")
    st.session_state.setdefault("at_tshr", "")
    st.session_state.setdefault("symp", "")
    for k in [
        "qa_family","qa_autoimmune","qa_postpartum","qa_radiation",
        "qa_iodine_low","qa_amiodarone","qa_lithium"
    ]:
        st.session_state.setdefault(k, False)
    st.session_state.setdefault("current_report", None)
    st.session_state.setdefault("history", load_json(HISTORY_JSON, []))
    st.session_state.setdefault("context_mode", "Обычный")
    st.session_state.setdefault("tirads_nodes", [])

def _hex_to_rgb(h: str):
    try:
        h = h.strip()
        if h.startswith("#"):
            h = h[1:]
        if len(h) == 3:
            h = "".join(ch*2 for ch in h)
        r = int(h[0:2], 16)
        g = int(h[2:4], 16)
        b = int(h[4:6], 16)
        return (r, g, b)
    except Exception:
        return (43, 108, 223)

def _rel_luminance(h: str) -> float:
    r, g, b = _hex_to_rgb(h)
    def _c(u):
        u = u / 255.0
        return u/12.92 if u <= 0.04045 else ((u+0.055)/1.055)**2.4
    R, G, B = _c(r), _c(g), _c(b)
    return 0.2126*R + 0.7152*G + 0.0722*B

def _contrast_ratio(c1: str, c2: str) -> float:
    L1 = _rel_luminance(c1)
    L2 = _rel_luminance(c2)
    Lb, Ld = (max(L1, L2), min(L1, L2))
    return (Lb + 0.05) / (Ld + 0.05)

def _resolve_title_color(settings: Dict[str, Any]) -> str:
    theme = settings.get("theme", "dark")
    title = settings.get("title_color") or ("#2b6cdf" if theme == "light" else "#eef2f7")
    bg = "#0e1117" if theme == "dark" else "#f7f7f9"
    if _contrast_ratio(title, bg) < 3.0:
        return "#2b6cdf" if theme == "light" else "#eef2f7"
    return title

def inject_style(accent: str, title_color: str, theme: str):
    bg = "#0e1117" if theme == "dark" else "#f7f7f9"
    card_bg = "rgba(255,255,255,0.08)" if theme == "dark" else "rgba(255,255,255,0.85)"
    text = "#eef2f7" if theme == "dark" else "#101318"
    st.markdown(
        f"""
        <style>
        :root {{
            --mnd-accent: {accent};
            --mnd-title: {title_color};
            --mnd-bg: {bg};
            --mnd-card: {card_bg};
            --mnd-text: {text};
        }}
        .stApp {{ background: var(--mnd-bg); color: var(--mnd-text); }}
        .mnd-header {{
            background: rgba(0,0,0,0.15);
            backdrop-filter: blur(14px);
            -webkit-backdrop-filter: blur(14px);
            border: 1px solid rgba(255,255,255,0.08);
            border-radius: 18px;
            padding: 18px 22px;
            margin: 8px 0 18px 0;
            color: var(--mnd-title);
            font-weight: 800;
            font-size: 30px;
        }}
        .mnd-underline {{
            height: 6px; width: 180px; border-radius: 8px;
            background: linear-gradient(90deg, var(--mnd-accent), {accent}80);
            margin: 6px 0 22px 0;
        }}
        .mnd-card {{
            background: var(--mnd-card);
            backdrop-filter: blur(16px);
            -webkit-backdrop-filter: blur(16px);
            border: 1px solid rgba(255,255,255,0.08);
            border-radius: 16px;
            padding: 18px;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )

def get_effective_refs(context_mode: str) -> Dict[str, float]:
    if context_mode.startswith("Беременность"):
        tr = context_mode.split()[-1]
        if tr == "I":
            return {"TSH_LOW": 0.1, "TSH_HIGH": 2.5, "FT4_LOW": 9.0, "FT4_HIGH": 19.0, "FT3_LOW": 3.1, "FT3_HIGH": 6.3}
        if tr == "II":
            return {"TSH_LOW": 0.2, "TSH_HIGH": 3.0, "FT4_LOW": 9.0, "FT4_HIGH": 19.0, "FT3_LOW": 3.1, "FT3_HIGH": 6.3}
        if tr == "III":
            return {"TSH_LOW": 0.3, "TSH_HIGH": 3.5, "FT4_LOW": 9.0, "FT4_HIGH": 19.0, "FT3_LOW": 3.1, "FT3_HIGH": 6.3}
    if context_mode.startswith("Педиатрия"):
        return {"TSH_LOW": 0.7, "TSH_HIGH": 6.0, "FT4_LOW": 10.0, "FT4_HIGH": 22.0, "FT3_LOW": 3.5, "FT3_HIGH": 6.5}
    return {"TSH_LOW": 0.4, "TSH_HIGH": 4.0, "FT4_LOW": 9.0, "FT4_HIGH": 19.0, "FT3_LOW": 3.5, "FT3_HIGH": 6.5}

WEIGHTS = {
    "hyper": {
        "tsh_low": 4, "ft4_high": 3, "ft3_high": 2,
        "symptom_hyper": 3, "at_tshr_pos": 3,
        "family": 2, "iodine_low": 2, "postpartum": 2,
    },
    "hypo": {"tsh_high": 4, "ft4_low": 3, "symptom_hypo": 2, "autoimmune": 2},
    "toxic_mn": {"at_tshr_pos": 4, "tsh_low": 3, "ft4_high": 2, "ft3_high": 2},
    "subclinical": {"tsh_abn": 3, "no_symptoms": 2},
    "thyroiditis": {"postpartum": 3, "autoimmune": 2, "tsh_abn": 2},
}

def parse_symptoms(symp_text: str) -> Dict[str, bool]:
    s = (symp_text or "").lower()
    hyper_hits = any(k in s for k in ["тахикард", "тремор", "потлив", "похуд", "нервоз"])
    hypo_hits  = any(k in s for k in ["вял", "сонл", "набор веса", "запор", "озноб", "сухость кожи"])
    return {"hyper": hyper_hits, "hypo": hypo_hits, "has_any": bool(s.strip())}

def compute_scores(tsh, ft4, ft3, at_tpo, at_tshr, symptoms_text, qa_flags: Dict[str, bool], refs: Dict[str, float]):
    sym = parse_symptoms(symptoms_text)
    tsh_low  = tsh is not None and tsh < refs["TSH_LOW"]
    tsh_high = tsh is not None and tsh > refs["TSH_HIGH"]
    ft4_low  = ft4 is not None and ft4 < refs["FT4_LOW"]
    ft4_high = ft4 is not None and ft4 > refs["FT4_HIGH"]
    ft3_low  = ft3 is not None and ft3 < refs["FT3_LOW"]
    ft3_high = ft3 is not None and ft3 > refs["FT3_HIGH"]
    at_tshr_pos = at_tshr is not None and at_tshr > 1.75

    def sum_w(keys: Dict[str,int], flags: Dict[str,bool]) -> int:
        return sum(w for k, w in keys.items() if flags.get(k, False))

    flags_hyper = {
        "tsh_low": tsh_low, "ft4_high": ft4_high, "ft3_high": ft3_high,
        "symptom_hyper": sym["hyper"], "at_tshr_pos": at_tshr_pos,
        "family": qa_flags.get("family", False),
        "iodine_low": qa_flags.get("iodine_low", False),
        "postpartum": qa_flags.get("postpartum", False),
    }
    flags_hypo = {
        "tsh_high": tsh_high, "ft4_low": ft4_low,
        "symptom_hypo": sym["hypo"], "autoimmune": qa_flags.get("autoimmune", False),
    }
    flags_toxic = {"at_tshr_pos": at_tshr_pos, "tsh_low": tsh_low, "ft4_high": ft4_high, "ft3_high": ft3_high}
    flags_subc  = {"tsh_abn": (tsh_low or tsh_high) and not (ft4_low or ft4_high or ft3_low or ft3_high),
                   "no_symptoms": not sym["has_any"]}
    flags_thyr  = {"postpartum": qa_flags.get("postpartum", False),
                   "autoimmune": qa_flags.get("autoimmune", False),
                   "tsh_abn": (tsh_low or tsh_high)}

    h_score = sum_w(WEIGHTS["hyper"], flags_hyper)
    p_score = sum_w(WEIGHTS["hypo"],  flags_hypo)
    t_score = sum_w(WEIGHTS["toxic_mn"], flags_toxic)
    s_score = sum_w(WEIGHTS["subclinical"], flags_subc)
    y_score = sum_w(WEIGHTS["thyroiditis"], flags_thyr)

    MAX_HYPER = max(1, sum(WEIGHTS["hyper"].values()))
    MAX_HYPO  = max(1, sum(WEIGHTS["hypo"].values()))
    MAX_TOX   = max(1, sum(WEIGHTS["toxic_mn"].values()))
    MAX_SUB   = max(1, sum(WEIGHTS["subclinical"].values()))
    MAX_THYR  = max(1, sum(WEIGHTS["thyroiditis"].values()))

    details = {
        "hyper": {"name": "Гипертиреоз", "pct": clamp_int(100.0 * h_score / MAX_HYPER)},
        "hypo":  {"name": "Гипотиреоз",  "pct": clamp_int(100.0 * p_score / MAX_HYPO)},
        "toxic_mn": {"name": "Токсический зоб", "pct": clamp_int(100.0 * t_score / MAX_TOX)},
        "subclinical": {"name": "Субклиническое", "pct": clamp_int(100.0 * s_score / MAX_SUB)},
        "thyroiditis": {"name": "Тиреоидит", "pct": clamp_int(100.0 * y_score / MAX_THYR)},
    }
    pk = max(details, key=lambda k: details[k]["pct"])
    primary = {"name": details[pk]["name"], "prob": details[pk]["pct"]}

    reasons: List[str] = []
    if tsh_low: reasons.append("ТТГ понижен")
    if tsh_high: reasons.append("ТТГ повышен")
    if ft4_high: reasons.append("FT4 повышен")
    if ft4_low: reasons.append("FT4 понижен")
    if ft3_high: reasons.append("FT3 повышен")
    if ft3_low: reasons.append("FT3 понижен")
    if at_tshr_pos: reasons.append("АТ-TSHR положителен")
    if sym["hyper"]: reasons.append("гипер-симптомы")
    if sym["hypo"]: reasons.append("гипо-симптомы")

    return details, primary, reasons

def build_report(pid, tsh, ft4, ft3, at_tpo, at_tshr, symptoms_text, qa: Dict[str, bool], refs: Dict[str, float]):
    details, primary, reasons = compute_scores(tsh, ft4, ft3, at_tpo, at_tshr, symptoms_text, qa, refs)
    return {
        "patient": pid or "—",
        "ts": datetime.now().strftime("%d.%m.%Y %H:%M"),
        "primary": primary,
        "details": details,
        "inputs": {"TSH": tsh, "FT4": ft4, "FT3": ft3, "AT_TPO": at_tpo, "AT_TSHR": at_tshr, "symptoms": symptoms_text or ""},
        "qa": qa,
        "reasons": reasons,
        "context": st.session_state.get("context_mode","Обычный"),
        "tirads": st.session_state.get("tirads_nodes", []),
    }

def render_probability_charts(details: Dict[str, Dict[str, Any]]):
    cats = ["Гипертиреоз", "Гипотиреоз", "Субклиническое", "Тиреоидит", "Токсический зоб"]
    vals = [details["hyper"]["pct"], details["hypo"]["pct"], details["subclinical"]["pct"], details["thyroiditis"]["pct"], details["toxic_mn"]["pct"]]
    df = pd.DataFrame({"Категория": cats, "Вероятность": vals})
    color_scale = alt.Scale(domain=cats, scheme='category10')
    donut = (alt.Chart(df)
             .mark_arc(innerRadius=80)
             .encode(
                 theta=alt.Theta("Вероятность:Q", stack=True),
                 color=alt.Color("Категория:N", scale=color_scale,
                                 legend=alt.Legend(orient="bottom", columns=3, labelLimit=1000, title=None)),
                 tooltip=[alt.Tooltip("Категория:N"), alt.Tooltip("Вероятность:Q")],
             ).properties(height=360))
    bars = (alt.Chart(df.sort_values("Вероятность", ascending=True))
            .mark_bar()
            .encode(
                x=alt.X("Вероятность:Q", title="%", axis=alt.Axis(labelLimit=1000)),
                y=alt.Y("Категория:N", sort="-x", axis=alt.Axis(labelLimit=1000)),
                color=alt.Color("Категория:N", scale=color_scale, legend=None),
                tooltip=[alt.Tooltip("Категория:N"), alt.Tooltip("Вероятность:Q")],
            ).properties(height=260))
    return donut, bars

def suggest_mkb(primary_name: str) -> List[str]:
    name = primary_name.lower()
    if "гипертиреоз" in name or "токсическ" in name:
        return ["E05", "E05.0–E05.9"]
    if "гипотиреоз" in name:
        return ["E03", "E03.0–E03.9"]
    if "тиреоидит" in name:
        return ["E06", "E06.0–E06.9"]
    return []

def parse_lis_text(txt: str) -> Dict[str, Optional[float]]:
    t = (txt or "").lower().replace(",", ".")
    import re
    def find_val(patterns: List[str]) -> Optional[float]:
        for p in patterns:
            m = re.search(p, t, re.IGNORECASE)
            if m:
                try:
                    return float(m.group(1))
                except Exception:
                    continue
        return None
    tsh = find_val([r"tsh[^0-9\-]*([\-]?\d+(\.\d+)?)", r"ттг[^0-9\-]*([\-]?\d+(\.\d+)?)", r"thyroid[- ]?stimulating[^0-9]*([\-]?\d+(\.\d+)?)"])
    ft4 = find_val([r"ft4[^0-9\-]*([\-]?\d+(\.\d+)?)", r"св[оё]б[оё]дн[ыйое]*\s*t4[^0-9\-]*([\-]?\d+(\.\d+)?)"])
    ft3 = find_val([r"ft3[^0-9\-]*([\-]?\d+(\.\d+)?)", r"св[оё]б[оё]дн[ыйое]*\s*t3[^0-9\-]*([\-]?\d+(\.\d+)?)"])
    return {"TSH": tsh, "FT4": ft4, "FT3": ft3}

def tirads_points(composition: str, echogenicity: str, shape_ttw: bool, margin: str, foci: str) -> Tuple[int, str, str]:
    comp_pts = {"Кистозный/спонг.": 0, "Смешанный": 1, "Солидный": 2}.get(composition, 0)
    echo_pts = {"Анехо/гиперэхо": 0, "Изоэхогенный": 1, "Гипоэхогенный": 2, "Выраженно гипоэхогенный": 3}.get(echogenicity, 0)
    shape_pts = 3 if shape_ttw else 0
    margin_pts = {"Гладкие": 0, "Неровные/дольчатые": 2, "Экстратиреоидные": 3}.get(margin, 0)
    foci_pts = {"Нет": 0, "Кометный хвост": 0, "Макрокальцинаты": 1, "Периферические": 2, "Микрокальцинаты": 3}.get(foci, 0)
    total = comp_pts + echo_pts + shape_pts + margin_pts + foci_pts
    if total <= 1: cat = "TR1"
    elif total == 2: cat = "TR2"
    elif 3 <= total <= 4: cat = "TR3"
    elif 5 <= total <= 6: cat = "TR4"
    else: cat = "TR5"
    if cat == "TR3":
        rec = "ФНА ≥2.5 см; наблюдение ≥1.5 см"
    elif cat == "TR4":
        rec = "ФНА ≥1.5 см; наблюдение ≥1.0 см"
    elif cat == "TR5":
        rec = "ФНА ≥1.0 см; наблюдение ≥0.5 см"
    else:
        rec = "Наблюдение по клинике/УЗИ"
    return total, cat, rec

def render_analysis():
    S = st.session_state.settings
    st.markdown(f'<div class="mnd-header">{APP_TITLE}</div>', unsafe_allow_html=True)
    st.markdown('<div class="mnd-underline"></div>', unsafe_allow_html=True)
    st.caption("Демонстрационный прототип. Не является медицинским изделием.")

    tcol1, tcol2 = st.columns([0.6, 0.4])
    with tcol1:
        if st.button("Пример данных"):
            st.session_state.update({
                "pid":"test_001","tsh":"0.05","ft4":"40","ft3":"8","at_tpo":"10","at_tshr":"2.0","symp":"тахикардия, тремор",
                "qa_family":False,"qa_autoimmune":False,"qa_postpartum":False,"qa_radiation":False,
                "qa_iodine_low":False,"qa_amiodarone":False,"qa_lithium":False,
            })
            st.rerun()
    with tcol2:
        st.selectbox("Клинический контекст", ["Обычный","Беременность I","Беременность II","Беременность III","Педиатрия (<18)"], key="context_mode")

    st.subheader("Ввод данных")
    st.text_input("ID пациента (опционально)", key="pid", placeholder="например, test_001")
    c1, c2 = st.columns(2)
    with c1:
        st.text_input("ТТГ (мЕд/л) *", key="tsh", placeholder="например, 1.2")
        st.text_input("FT4 (пмоль/л)", key="ft4", placeholder="например, 12")
        st.text_input("FT3 (пмоль/л)", key="ft3", placeholder="например, 4.5")
    with c2:
        st.text_input("АТ-ТПО (Ед/мл)", key="at_tpo", placeholder="например, 10")
        st.text_input("АТ-TSHR (ед.)", key="at_tshr", placeholder="например, 1.5")
        st.text_input("Симптомы (через запятую)", key="symp", placeholder="например, тахикардия, тремор")

    with st.expander("Быстрый ввод из ЛИС (вставьте текст и нажмите «Распознать»)", expanded=False):
        lis_txt = st.text_area("Вставьте сюда текст лабораторной панели", height=120, placeholder="TSH 0.05 мЕд/л; FT4 40 пмоль/л; FT3 8 пмоль/л ...")
        if st.button("Распознать из текста"):
            parsed = parse_lis_text(lis_txt)
            if parsed["TSH"] is not None: st.session_state.tsh = str(parsed["TSH"])
            if parsed["FT4"] is not None: st.session_state.ft4 = str(parsed["FT4"])
            if parsed["FT3"] is not None: st.session_state.ft3 = str(parsed["FT3"])
            st.success("Значения распознаны и подставлены.")
            st.rerun()

    st.markdown("### Факторы риска (опрос):")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.checkbox("Семейный анамнез", key="qa_family")
        st.checkbox("Аутоиммунные заболевания", key="qa_autoimmune")
        st.checkbox("Послеродовой период", key="qa_postpartum")
    with c2:
        st.checkbox("Облучение/область шеи", key="qa_radiation")
        st.checkbox("Дефицит йода", key="qa_iodine_low")
    with c3:
        st.checkbox("Амиодарон", key="qa_amiodarone")
        st.checkbox("Литий", key="qa_lithium")

    with st.expander("УЗИ: узлы TI-RADS", expanded=False):
        ncol1, ncol2, ncol3 = st.columns(3)
        with ncol1:
            size_a = st.number_input("Размер A (мм)", min_value=0.0, step=0.1, value=0.0, key="n_a")
            size_b = st.number_input("Размер B (мм)", min_value=0.0, step=0.1, value=0.0, key="n_b")
            size_c = st.number_input("Размер C (мм)", min_value=0.0, step=0.1, value=0.0, key="n_c")
        with ncol2:
            composition = st.selectbox("Состав", ["Кистозный/спонг.","Смешанный","Солидный"], key="n_comp")
            echogenicity = st.selectbox("Эхогенность", ["Анехо/гиперэхо","Изоэхогенный","Гипоэхогенный","Выраженно гипоэхогенный"], key="n_echo")
            margin = st.selectbox("Контуры", ["Гладкие","Неровные/дольчатые","Экстратиреоидные"], key="n_margin")
        with ncol3:
            shape_ttw = st.checkbox("Выше, чем шире (taller-than-wide)", key="n_ttw")
            foci = st.selectbox("Эхогенные включения", ["Нет","Кометный хвост","Макрокальцинаты","Периферические","Микрокальцинаты"], key="n_foci")
            if st.button("Добавить узел"):
                points, cat, rec = tirads_points(composition, echogenicity, shape_ttw, margin, foci)
                node = {
                    "a_mm": size_a, "b_mm": size_b, "c_mm": size_c,
                    "comp": composition, "echo": echogenicity, "margin": margin,
                    "ttw": bool(shape_ttw), "foci": foci,
                    "points": points, "cat": cat, "rec": rec
                }
                arr = list(st.session_state.tirads_nodes or [])
                arr.append(node)
                st.session_state.tirads_nodes = arr
                st.success(f"Узел добавлен: {cat}, {points} баллов.")
                st.rerun()
        nodes = st.session_state.tirads_nodes or []
        if nodes:
            df_nodes = pd.DataFrame(nodes)
            st.dataframe(df_nodes, use_container_width=True, hide_index=True)
            if st.button("Очистить список узлов"):
                st.session_state.tirads_nodes = []
                st.rerun()

    run_full = st.button("Анализировать", type="primary", use_container_width=True)
    if run_full:
        tsh = safe_float(st.session_state.tsh)
        if tsh is None:
            st.error("Введите корректный ТТГ")
        else:
            refs = get_effective_refs(st.session_state.context_mode)
            report = build_report(
                st.session_state.pid,
                tsh,
                safe_float(st.session_state.ft4),
                safe_float(st.session_state.ft3),
                safe_float(st.session_state.at_tpo),
                safe_float(st.session_state.at_tshr),
                st.session_state.symp,
                {
                    "family": bool(st.session_state.qa_family),
                    "autoimmune": bool(st.session_state.qa_autoimmune),
                    "postpartum": bool(st.session_state.qa_postpartum),
                    "radiation": bool(st.session_state.qa_radiation),
                    "iodine_low": bool(st.session_state.qa_iodine_low),
                    "amiodarone": bool(st.session_state.qa_amiodarone),
                    "lithium": bool(st.session_state.qa_lithium),
                },
                refs
            )
            st.session_state.current_report = report
            if st.session_state.settings.get("autosave_history", True):
                hist = load_json(HISTORY_JSON, [])
                hist.append(report)
                save_json(HISTORY_JSON, hist)
                st.session_state.history = hist

    st.markdown("## Результаты")
    if st.session_state.current_report:
        r = st.session_state.current_report
        def fmt(v): return "—" if v is None or v == "" else v
        def rus_bool(b): return "Да" if b else "Нет"

        risk_rows = [
            ("Семейный анамнез", rus_bool(r['qa']['family'])),
            ("Аутоиммунные заболевания", rus_bool(r['qa']['autoimmune'])),
            ("Послеродовой период", rus_bool(r['qa']['postpartum'])),
            ("Облучение/область шеи", rus_bool(r['qa']['radiation'])),
            ("Дефицит йода", rus_bool(r['qa']['iodine_low'])),
            ("Амиодарон", rus_bool(r['qa']['amiodarone'])),
            ("Литий", rus_bool(r['qa']['lithium'])),
        ]
        risks_html = "".join([f"<li>{k}: {v}</li>" for k, v in risk_rows])

        tirads_block = ""
        if r.get("tirads"):
            lines = []
            for i, n in enumerate(r["tirads"], start=1):
                dmax = max(n.get("a_mm",0), n.get("b_mm",0), n.get("c_mm",0))
                lines.append(f"Узел {i}: {n['cat']} ({n['points']} б.), макс. размер ~{dmax:.1f} мм; {n['rec']}")
            tirads_block = "<div><b>УЗИ (TI-RADS):</b><ul>" + "".join([f"<li>{x}</li>" for x in lines]) + "</ul></div>"

        mkb_list = suggest_mkb(r["primary"]["name"])
        mkb_html = f"<div><b>МКБ-10 (ориентировочно):</b> {', '.join(mkb_list) if mkb_list else '—'}</div>"

        conclusion_html = f"""
        <div class="mnd-card">
          <div style="font-weight:700;margin-bottom:12px;">Клиническое заключение</div>
          <div style="opacity:.85">
            <b>Пациент:</b> {r['patient']}<br/>
            <b>Дата/время:</b> {r['ts']}<br/>
            <b>Контекст:</b> {r.get('context','Обычный')}<br/>
            <b>Основная вероятность:</b> {r['primary']['name']} — ~{r['primary']['prob']}%
          </div>
          <div style="height:8px"></div>
          <div>
            <b>Входные данные:</b>
            <ul>
              <li>ТТГ: {fmt(r['inputs']['TSH'])}</li>
              <li>FT4: {fmt(r['inputs']['FT4'])}</li>
              <li>FT3: {fmt(r['inputs']['FT3'])}</li>
              <li>АТ-ТПО: {fmt(r['inputs']['AT_TPO'])}</li>
              <li>АТ-TSHR: {fmt(r['inputs']['AT_TSHR'])}</li>
              <li>Симптомы: {fmt(r['inputs']['symptoms'])}</li>
            </ul>
          </div>
          <div><b>Факторы риска:</b>
            <ul style="columns:2;-webkit-columns:2;-moz-columns:2;">{risks_html}</ul>
          </div>
          {tirads_block}
          {mkb_html}
        </div>
        """
        st.markdown(conclusion_html, unsafe_allow_html=True)

        if S.get("embed_chart", True):
            donut, bars = render_probability_charts(r["details"])
            c1, c2 = st.columns(2)
            with c1:
                st.subheader("Круговая диаграмма")
                st.altair_chart(donut, use_container_width=True)
            with c2:
                st.subheader("Горизонтальная диаграмма")
                st.altair_chart(bars, use_container_width=True)

        if st.session_state.history:
            try:
                df_hist = pd.DataFrame([
                    {
                        "Пациент": h.get("patient", "—"),
                        "Дата": h.get("ts", ""),
                        "TSH": h.get("inputs", {}).get("TSH"),
                        "FT4": h.get("inputs", {}).get("FT4"),
                        "FT3": h.get("inputs", {}).get("FT3"),
                    }
                    for h in st.session_state.history
                ])

                pid = (st.session_state.pid or "").strip()
                if pid:
                    df_hist = df_hist[df_hist["Пациент"] == pid]

                if not df_hist.empty:
                    df_hist["Дата_dt"] = pd.to_datetime(
                        df_hist["Дата"], format="%d.%m.%Y %H:%M", errors="coerce"
                    )
                    df_hist = df_hist.dropna(subset=["Дата_dt"]).sort_values("Дата_dt")

                    if not df_hist.empty:
                        st.subheader("Динамика показателей")
                        base = alt.Chart(df_hist).encode(
                            x=alt.X(
                                "Дата_dt:T",
                                axis=alt.Axis(format="%d.%m.%y", labelAngle=0, labelLimit=120),
                                title="Дата",
                            )
                        )
                        st.altair_chart(
                            base.mark_line().encode(
                                y=alt.Y("TSH:Q", title="ТТГ"),
                                tooltip=["Дата", "TSH"]
                            ),
                            use_container_width=True,
                        )
                        st.altair_chart(
                            base.mark_line().encode(
                                y=alt.Y("FT4:Q", title="FT4"),
                                tooltip=["Дата", "FT4"]
                            ),
                            use_container_width=True,
                        )
                        st.altair_chart(
                            base.mark_line().encode(
                                y=alt.Y("FT3:Q", title="FT3"),
                                tooltip=["Дата", "FT3"]
                            ),
                            use_container_width=True,
                        )
            except Exception as e:
                st.caption(f"Не удалось построить графики: {e}")

        txt_lines = [
            f"Пациент: {r['patient']}",
            f"Дата/время: {r['ts']}",
            f"Контекст: {r.get('context','Обычный')}",
            f"Основная вероятность: {r['primary']['name']} — ~{r['primary']['prob']}%",
            "",
            "Входные данные:",
            f"  - TSH: {fmt(r['inputs']['TSH'])}",
            f"  - FT4: {fmt(r['inputs']['FT4'])}",
            f"  - FT3: {fmt(r['inputs']['FT3'])}",
            f"  - AT_TPO: {fmt(r['inputs']['AT_TPO'])}",
            f"  - AT_TSHR: {fmt(r['inputs']['AT_TSHR'])}",
            f"  - symptoms: {fmt(r['inputs']['symptoms'])}",
            "",
            "Факторы риска:"
        ] + [f"  - {k}: {v}" for k, v in [
            ("family", rus_bool(r['qa']['family'])),
            ("autoimmune", rus_bool(r['qa']['autoimmune'])),
            ("postpartum", rus_bool(r['qa']['postpartum'])),
            ("radiation", rus_bool(r['qa']['radiation'])),
            ("iodine_low", rus_bool(r['qa']['iodine_low'])),
            ("amiodarone", rus_bool(r['qa']['amiodarone'])),
            ("lithium", rus_bool(r['qa']['lithium'])),
        ]]

        if r.get("tirads"):
            txt_lines += ["", "УЗИ (TI-RADS):"]
            for i, n in enumerate(r["tirads"], start=1):
                dmax = max(n.get("a_mm",0), n.get("b_mm",0), n.get("c_mm",0))
                txt_lines.append(f"  - Узел {i}: {n['cat']} ({n['points']} б.), макс. размер ~{dmax:.1f} мм; {n['rec']}")

        mkb = suggest_mkb(r["primary"]["name"])
        if mkb:
            txt_lines += ["", "МКБ-10 (ориентировочно): " + ", ".join(mkb)]

        final_text = "\n".join(txt_lines)

        st.markdown("### Экспорт заключения")
        st.text_area("Текст заключения", value=final_text, height=180)
        st.download_button("Скачать .txt", data=final_text.encode("utf-8-sig"), file_name="mnd_report.txt", mime="text/plain")

        if Document is not None:
            if st.button("Скачать DOCX"):
                buf = io.BytesIO()
                doc = Document()
                p = doc.add_paragraph(APP_TITLE)
                p.runs[0].font.size = Pt(14)
                for line in txt_lines:
                    doc.add_paragraph(line)
                doc.save(buf)
                st.download_button("Загрузить файл DOCX", data=buf.getvalue(), file_name="mnd_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.caption("Для экспорта в DOCX установите пакет: pip install python-docx")
    else:
        st.info("Заполните поля и нажмите «Анализировать».")

def render_sources():
    st.markdown(f'<div class="mnd-header">Источники</div>', unsafe_allow_html=True)
    st.markdown('<div class="mnd-underline"></div>', unsafe_allow_html=True)
    sources = [
        ("СИБАК: материалы конференции (медицина)", "Публикация по тематике эндокринологии/ЩЖ (RU).", "https://sibac.info/conf/med/ii/26124"),
        ("«Проблемы эндокринологии» — статья", "Публикация по эндокринологии (RU).", "https://endocrinology-journal.ru/ru/jarticles_endo/841.html?SSr=07E90819114448"),
        ("«Проблемы эндокринологии» — статья", "Публикация по эндокринологии (RU).", "https://endocrinology-journal.ru/ru/jarticles_endo/498.html?SSr=07E90819114448"),
        ("«Проблемы эндокринологии» — статья", "Публикация по эндокринологии (RU).", "https://endocrinology-journal.ru/ru/jarticles_endo/537.html?SSr=07E90819114448"),
        ("«Проблемы эндокринологии» — официальный сайт", "Журнал НМИЦ эндокринологии (RU).", "https://endocrinology-journal.ru/"),
        ("NICE NG145: Болезни щитовидной железы — оценка и ведение", "Официальное руководство NICE (PDF).", "https://www.entuk.org/_userfiles/pages/files/guidelines/thyroid-disease-assessment-and-management-pdf-66141781496773.pdf"),
        ("ATA 2016: Тиреотоксикоз — клинические рекомендации", "American Thyroid Association, 2016 (PubMed).", "https://pubmed.ncbi.nlm.nih.gov/27521067/"),
        ("ATA 2014: Лечение гипотиреоза — рекомендации", "American Thyroid Association, 2014 (PubMed).", "https://pubmed.ncbi.nlm.nih.gov/25266247/"),
        ("AACE/ACE 2012: Гипотиреоз у взрослых — рекомендации", "American Association of Clinical Endocrinologists, 2012 (PubMed).", "https://pubmed.ncbi.nlm.nih.gov/23246686/"),
        ("ETA 2018: Болезнь Грейвса — рекомендации", "European Thyroid Association, 2018 (PubMed).", "https://pubmed.ncbi.nlm.nih.gov/30283735/"),
        ("ETA 2018: Центральный гипотиреоз — рекомендации", "European Thyroid Association, 2018 (PubMed).", "https://pubmed.ncbi.nlm.nih.gov/30374425/"),
        ("ETA 2018: Амидарон-ассоциированные нарушения ЩЖ — рекомендации", "European Thyroid Association, 2018 (PubMed).", "https://pubmed.ncbi.nlm.nih.gov/29594056/"),
        ("Послеродовой тиреоидит — обзор", "Обзор, 2019 (PubMed).", "https://pubmed.ncbi.nlm.nih.gov/30844908/"),
        ("Аутоиммунный тиреоидит Хашимото — обзор", "Обзор, 2020 (PubMed).", "https://pubmed.ncbi.nlm.nih.gov/32805423/"),
        ("Дефицит йода — современный обзор", "Endocrine Reviews, 2022 (PMC).", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9459956/"),
        ("КиберЛенинка", "Русскоязычные статьи по эндокринологии.", "https://cyberleninka.ru/"),
        ("American Thyroid Association (ATA)", "Образовательные материалы и клиническая информация.", "https://www.thyroid.org/"),
        ("Endocrine Society — Clinical Practice Guidelines", "Клинические руководства по эндокринологии.", "https://www.endocrine.org/clinical-practice-guidelines"),
    ]
    for name, desc, url in sources:
        cont = st.container(border=True)
        cols = cont.columns([0.75, 0.25])
        with cols[0]:
            st.markdown(f"**{name}**  \n{desc}")
            st.caption(url)
        with cols[1]:
            st.link_button("Открыть", url, use_container_width=True)

def render_history_export():
    st.markdown(f'<div class="mnd-header">История / Экспорт</div>', unsafe_allow_html=True)
    st.markdown('<div class="mnd-underline"></div>', unsafe_allow_html=True)
    hist = st.session_state.history
    if not hist:
        st.info("История пуста.")
        return
    rows = [{"Пациент": r.get("patient","—"),
             "Дата/время": r.get("ts",""),
             "Основная вероятность": f"{r['primary']['name']} ~{r['primary']['prob']}%"} for r in hist]
    df = pd.DataFrame(rows)
    st.dataframe(df, use_container_width=True, hide_index=True)
    csv = df.to_csv(index=False).encode("utf-8-sig")
    st.download_button("Скачать CSV", data=csv, file_name="mnd_history.csv", mime="text/csv")
    if st.button("Очистить историю", type="secondary"):
        save_json(HISTORY_JSON, [])
        st.session_state.history = []
        st.success("История очищена.")

def render_settings():
    S = st.session_state.settings
    st.markdown(f'<div class="mnd-header">Настройки</div>', unsafe_allow_html=True)
    st.markdown('<div class="mnd-underline"></div>', unsafe_allow_html=True)
    theme = st.radio("Тема приложения", ["Тёмная","Светлая"],
                     index=0 if S.get("theme","dark")=="dark" else 1, horizontal=True)
    theme_val = "dark" if theme=="Тёмная" else "light"
    accent = st.color_picker("Акцент (верх/кнопки)", value=S.get("accent_color","#2b6cdf"))
    title_color_input = st.color_picker("Цвет названия", value=S.get("title_color","#2b6cdf"))
    embed = st.checkbox("Встраивать диаграмму справа", value=S.get("embed_chart", True))
    autosave = st.checkbox("Автосохранение истории", value=S.get("autosave_history", True))
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Применить", use_container_width=True):
            tmp = {"theme": theme_val, "accent_color": accent, "title_color": title_color_input,
                   "embed_chart": bool(embed), "autosave_history": bool(autosave)}
            tmp["title_color"] = _resolve_title_color(tmp)
            S.update(tmp)
            save_json(SETTINGS_JSON, S)
            st.success("Настройки применены.")
            st.rerun()
    with c2:
        if st.button("Сбросить к умолчанию", use_container_width=True):
            st.session_state.settings = DEFAULT_SETTINGS.copy()
            st.session_state.settings["title_color"] = _resolve_title_color(st.session_state.settings)
            save_json(SETTINGS_JSON, st.session_state.settings)
            st.success("Сброшено.")
            st.rerun()

def render_instructions():
    st.markdown(f'<div class="mnd-header">Инструкция</div>', unsafe_allow_html=True)
    st.markdown('<div class="mnd-underline"></div>', unsafe_allow_html=True)
    st.markdown(
        """
### Назначение
Учебный прототип для ориентировочной оценки риска нарушений функции щитовидной железы по лабораторным данным, симптомам и факторам риска. Решение врача не заменяет.

### Подготовка данных (единицы СИ)
- **ТТГ** — мЕд/л  
- **FT4** — пмоль/л  
- **FT3** — пмоль/л  
- **АТ-ТПО** — Ед/мл  
- **АТ-TSHR** — ед. метода (в прототипе положительным считается > 1.75)

Симптомы вводите через запятую: тахикардия, тремор.

### Порядок работы
1. Во вкладке **Анализ** заполните поля или нажмите **Пример данных**.  
2. Выберите **Клинический контекст** (Обычный / Беременность I–III / Педиатрия).  
3. Отметьте **Факторы риска**.  
4. При необходимости воспользуйтесь блоком **Быстрый ввод из ЛИС** — вставьте текст и нажмите «Распознать».  
5. Опционально заполните **УЗИ: TI-RADS** для узлов (размеры в мм, признаки) и добавьте узлы.  
6. Нажмите **Анализировать**.  
7. Ниже появится **Клиническое заключение**, диаграммы и при наличии узлов — блок TI-RADS.  
8. Во вкладке **История/Экспорт** можно скачать CSV, а в самом заключении — TXT / DOCX.

### Интерпретация результатов
- **Основная вероятность** — наиболее вероятное состояние из: Гипертиреоз / Гипотиреоз / Субклиническое / Тиреоидит / Токсический зоб.  
- Проценты — нормированный скоринг, учитывающий лабораторные показатели, симптомы и факторы риска.  
- Блок **TI-RADS** рассчитывает категорию TR1–TR5 и даёт пороговые рекомендации (наблюдение/ФНА) по размерам узла.

### Референсы, используемые в расчётах
- Обычный adultes: ТТГ 0.4–4.0 мЕд/л; FT4 9–19 пмоль/л; FT3 3.5–6.5 пмоль/л.  
- Беременность: автоматическая подстройка диапазона ТТГ по триместрам.  
- Педиатрия: расширенный диапазон ТТГ и FT4.

### МКБ-10 (ориентировочно)
- Гипертиреоз (в т.ч. токсический зоб) — **E05**, диапазон **E05.0–E05.9**  
- Гипотиреоз — **E03**, диапазон **E03.0–E03.9**  
- Тиреоидит — **E06**, диапазон **E06.0–E06.9**  

### Частые ошибки и решения
- **ТТГ пустой** → расчёт не стартует. Введите число (разделитель , или . — допустимы оба).  
- **Нет DOCX-кнопки** → установить pip install python-docx.  
- **Не сохраняются история/настройки** → нет прав записи в папку. Запустите из директории с правами.

### Ограничения
Прототип предназначен для демонстрации. Не является медицинским изделием. Окончательное решение — за лечащим врачом.
"""
    )

def render_about():
    st.markdown(f'<div class="mnd-header">О приложении</div>', unsafe_allow_html=True)
    st.markdown('<div class="mnd-underline"></div>', unsafe_allow_html=True)
    st.markdown(
        """
**МНД — Анализатор** — учебный прототип поддержки принятия решений для эндокринологии ЩЖ.
Логика основана на прозрачных правилах и взвешенном скоринге:

- Лабораторные показатели (**ТТГ, FT4, FT3, АТ-ТПО, АТ-TSHR**),  
- Симптомы (словари ключевых признаков гипер-/гипофункции),  
- Факторы риска (семейный анамнез, аутоиммунные, послеродовый период, радиация, дефицит йода, амиодарон, литий),  
- Контекст (беременность по триместрам, педиатрия) — корректирует «нормы» ТТГ/FT4/FT3,  
- УЗИ-узлы (TI-RADS): вычисление категории TR1–TR5 и пороговых рекомендаций наблюдения/ФНА.

Проценты по категориям — это **нормированный суммарный балл**: признаки имеют веса, сумма весов категории служит делителем, поэтому масштабы сопоставимы между собой.  
МКБ-10: гипертиреоз/токсический зоб **E05**, гипотиреоз **E03**, тиреоидит **E06** — **ориентиры**, а не автоматический диагноз.

**Адаптация под РФ**: единицы **СИ**, русскоязычная терминология, источники включают российские публикации и международные руководства (NICE, ATA/ETA, PubMed/PMC).

Прототип не заменяет клиническое мышление и не предназначен для самостоятельной диагностики без валидации.
"""
    )

def main():
    st.set_page_config(page_title=APP_TITLE, layout="wide", initial_sidebar_state="collapsed")
    ensure_state_defaults()
    S = st.session_state.settings
    effective_title = _resolve_title_color(S)
    inject_style(S.get("accent_color","#2b6cdf"), effective_title, S.get("theme","dark"))
    tabs = st.tabs(["Анализ", "Источники", "История/Экспорт", "Настройки", "Инструкция", "О приложении"])
    with tabs[0]: render_analysis()
    with tabs[1]: render_sources()
    with tabs[2]: render_history_export()
    with tabs[3]: render_settings()
    with tabs[4]: render_instructions()
    with tabs[5]: render_about()

if __name__ == "__main__":
    main()
